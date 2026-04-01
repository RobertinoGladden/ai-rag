from typing import List
from pathlib import Path
from loguru import logger

from .base_loader import BaseLoader, Document


class DocxLoader(BaseLoader):
    """Loader untuk file .docx menggunakan python-docx."""

    @property
    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]

    def load(self, source: str) -> List[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

        path = Path(source)
        if not path.exists():
            raise FileNotFoundError(f"File tidak ditemukan: {source}")

        logger.info(f"Loading DOCX: {path.name}")
        doc = DocxDocument(str(path))

        # Ambil semua paragraf yang tidak kosong
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        # Ambil teks dari tabel juga
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        if table_texts:
            content += "\n\n[Tables]\n" + "\n".join(table_texts)

        return [Document(
            content=content,
            metadata={
                "source": str(path),
                "filename": path.name,
                "type": "docx",
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
            }
        )]
